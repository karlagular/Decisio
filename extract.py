import os
from docx import Document

#define a function that returns the content
#def extract():

#define a function that converts the .doc files to .docx 
# (to be executed if the file ends in .doc)
#def doc2docx(doc):

#List al files in Porjectreports
path = 'Projectreports'
files = os.listdir(path)
#print(files)

# Filter the list to include only .docx files
docx_files = [file for file in files if file.endswith('.docx')]

# List of the names of the subtitles with style Heading 3 (only those in the first chapter)
headings_3=['Projektbeschreibung und -design', #1.1.1
          'Beschreibung des Projekterfolges aus Sicht der Kunden/des Auftraggebers' #1.1.2
          ] 

# List of the names of the subtitles with style Heading 2 (from chapter 2 untill chapter 5)
headings_2=['Steckbrief', #2.1
            'Ziele', #2.2
            'Priorisierung ausgewählter konkurrierender Ziele', #2.3
            'Abnahmekriterien', #3.1
            'Umfeldportfolio', #4.1
            'Stakeholder: Interessen, Erwartungen, Befürchtungen, Maßnahmen', #4.2
            'Stakeholderportfolio', #4.3
            'Erfassung und Benennung von drei Risiken', #5.1
            'Maßnahmen und Berechnung', #5.2
            'Erfassung und Benennung einer Chance' #5.3
            ]

# Create empty dictionary. Subchapter names are keys and the paragraph content is the value.
chapt_paragraph = {}

for chapt in headings_3:
    chapt_paragraph[chapt] = None
for chapt in headings_2:
    chapt_paragraph[chapt] = None

#print(chapt_paragraph)
#print("\n")


# The text of the chapter will be saved into this variable
textvar=""

if docx_files:
    # Open the first .docx file found in the directory
    document = Document(os.path.join(path, docx_files[0]))

    #flag initialised to false
    found_heading = False

    for p in document.paragraphs:
        if found_heading:
            # Extract text until the next Heading 3
            if p.style.name == 'Heading 3':
                break
            # variable filled with all paragraphs in current chapter
            textvar += f"{p.text}"
        # if the Subtitle matches the first item on the list of headings...
        elif p.style.name == 'Heading 3' and p.text == headings_3[0]:
            # we found the heading we're looking for
            found_heading = True
            #print(p.text)
        
    # Fill the first Dictionary value
    chapt_paragraph[headings_3[0]]=textvar
        
    # We repeat the same code again, can be optimised later on to be called a sa function 
    # with the arguments of what we're searching for and where to stop (condition for break)
    
    # flag reset to false 
    found_heading = False

    # textvar reset to empty string
    textvar = ''

    for p in document.paragraphs:
        if found_heading:
            # Extract text until the next Heading 2
            if p.style.name == 'Heading 1':
                break
            # variable filled with all paragraphs in current chapter
            textvar += f"{p.text}"
        # if the Subtitle matches the first item on the list of headings...
        elif p.style.name == 'Heading 3' and p.text == headings_3[1]:
            # we found the heading we're looking for
            found_heading = True
            #print(p.text)
    # Fill the first Dictionary value
    chapt_paragraph[headings_3[1]]=textvar   

    # Repeat again but for the rest of the chapters using headings_2 inside a loop
    for i in range(len(headings_2)):
        # flag reset to false 
        found_heading = False
        # textvar reset to empty string
        textvar = ''

        for p in document.paragraphs:
            if found_heading:
                # Extract text until the next Heading 2
                if p.style.name == 'Heading 2' or p.style.name=='Heading 1':
                    break
                # variable filled with all paragraphs in current chapter that aren't table descriptions
                if not p.text.startswith("Tabelle"):
                    textvar += f"{p.text}"
            # if the Subtitle matches the first item on the list of headings...
            elif p.style.name == 'Heading 2' and p.text == headings_2[i-1]:
                # we found the heading we're looking for
                found_heading = True
                #print(p.text)
        # Fill the first Dictionary value
        chapt_paragraph[headings_2[i-1]]=textvar
    
else:
    print("No .docx files found in the directory.")

#print(textvar)


print(chapt_paragraph)